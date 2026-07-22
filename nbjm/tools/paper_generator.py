import os
import sys
import json
import re
import subprocess
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

MARKDOCX_PATH = r"D:\D下载\markdocx-master\markdocx-master"

try:
    import latex2mathml.converter
    HAS_LATEX2MATHML = True
except ImportError:
    HAS_LATEX2MATHML = False

try:
    from lxml import etree
    HAS_LXML = True
except ImportError:
    HAS_LXML = False

def find_mml2omml_xsl():
    possible_paths = [
        r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL',
        r'C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL',
        r'C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL',
        r'C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL',
        r'C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL',
        r'C:\Program Files (x86)\Microsoft Office\root\Office15\MML2OMML.XSL',
        r'C:\Program Files\Microsoft Office\Office15\MML2OMML.XSL',
        r'C:\Program Files (x86)\Microsoft Office\Office15\MML2OMML.XSL',
        os.path.join(os.path.dirname(__file__), 'MML2OMML.XSL'),
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            return path
    return None

MML2OMML_XSL_PATH = find_mml2omml_xsl()

def sanitize_text(text):
    if text is None:
        return ""
    text = str(text)
    text = text.replace('\x00', '')
    text = text.replace('\n', ' ')
    text = text.replace('\r', ' ')
    text = text.replace('\t', ' ')
    text = re.sub(r'[\x00-\x1F\x7F]', '', text)
    return text.strip()

def preprocess_latex_formula(latex_formula):
    formula = latex_formula.strip()
    if formula.startswith('$$') and formula.endswith('$$'):
        formula = formula[2:-2]
    elif formula.startswith('$') and formula.endswith('$'):
        formula = formula[1:-1]
    
    formula = formula.strip()
    
    formula = formula.replace('\\mathbf', '\\boldsymbol')
    formula = formula.replace('\\mathrm', '')
    formula = formula.replace('\\text', '')
    formula = formula.replace('\\max', '\\mathrm{max}')
    formula = formula.replace('\\min', '\\mathrm{min}')
    formula = formula.replace('\\int', '\\int')
    
    return formula

def add_latex_formula(paragraph, latex_formula):
    try:
        cleaned_formula = latex_formula.strip()
        is_block = False
        if cleaned_formula.startswith('$$') and cleaned_formula.endswith('$$'):
            inner_formula = cleaned_formula[2:-2]
            is_block = True
        elif cleaned_formula.startswith('$') and cleaned_formula.endswith('$'):
            inner_formula = cleaned_formula[1:-1]
        else:
            inner_formula = cleaned_formula
        
        inner_formula = preprocess_latex_formula(inner_formula)
        
        if HAS_LATEX2MATHML and HAS_LXML and MML2OMML_XSL_PATH:
            try:
                mathml = latex2mathml.converter.convert(inner_formula)
                
                xslt = etree.parse(MML2OMML_XSL_PATH)
                transform = etree.XSLT(xslt)
                
                mathml_root = etree.fromstring(mathml)
                omml_root = transform(mathml_root)
                
                omml_string = etree.tostring(omml_root, pretty_print=False, encoding='unicode')
                
                run = paragraph.add_run()
                run._element.append(parse_xml(omml_string))
                
                return True
            except Exception as e:
                pass
        
        run = paragraph.add_run('$' + inner_formula + '$')
        run.font.name = 'Cambria Math'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria Math')
        run._element.rPr.rFonts.set(qn('w:cs'), 'Cambria Math')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Cambria Math')
        if is_block:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        return True
    except Exception as e:
        try:
            run = paragraph.add_run(latex_formula)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(12)
        except:
            paragraph.add_run(latex_formula)
        return False

def add_table_to_docx(doc, table_data):
    try:
        caption = sanitize_text(table_data.get('caption', ''))
        headers = [sanitize_text_preserve_newlines(h) for h in table_data.get('headers', [])]
        rows = [[sanitize_text_preserve_newlines(cell) for cell in row] for row in table_data.get('rows', [])]
        
        if caption:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(11)
            caption_run.bold = True
        
        if headers and rows:
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            for row_idx, row_data in enumerate(rows):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(row_cells):
                        cell = row_cells[col_idx]
                        if '$' in cell_data:
                            cell.paragraphs[0].text = ''
                            parts = split_content_with_formulas(cell_data)
                            for part in parts:
                                if part['type'] == 'text':
                                    cell.paragraphs[0].add_run(part['content'])
                                elif part['type'] == 'inline_formula':
                                    add_latex_formula(cell.paragraphs[0], part['content'])
                                elif part['type'] == 'block_formula':
                                    add_latex_formula(cell.paragraphs[0], part['content'])
                        else:
                            cell.text = cell_data
            
            table.style = 'Table Grid'
            
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
        
        doc.add_paragraph()
        return True
    except Exception as e:
        return False

def split_content_with_formulas(content):
    result = []
    pos = 0
    
    while pos < len(content):
        block_match = re.match(r'\$\$([\s\S]*?)\$\$', content[pos:])
        inline_match = re.match(r'\$([^$\n]+?)\$', content[pos:])
        
        if block_match and inline_match:
            if block_match.start() < inline_match.start():
                if pos < block_match.start():
                    text = content[pos:pos + block_match.start()]
                    if text.strip():
                        result.append({'type': 'text', 'content': text})
                formula_content = block_match.group(1)
                if formula_content.strip():
                    result.append({'type': 'block_formula', 'content': '$$' + formula_content + '$$'})
                pos += block_match.end()
            else:
                if pos < inline_match.start():
                    text = content[pos:pos + inline_match.start()]
                    if text.strip():
                        result.append({'type': 'text', 'content': text})
                formula_content = inline_match.group(1)
                if formula_content.strip():
                    result.append({'type': 'inline_formula', 'content': '$' + formula_content + '$'})
                pos += inline_match.end()
        elif block_match:
            if pos < block_match.start():
                text = content[pos:pos + block_match.start()]
                if text.strip():
                    result.append({'type': 'text', 'content': text})
            formula_content = block_match.group(1)
            if formula_content.strip():
                result.append({'type': 'block_formula', 'content': '$$' + formula_content + '$$'})
            pos += block_match.end()
        elif inline_match:
            if pos < inline_match.start():
                text = content[pos:pos + inline_match.start()]
                if text.strip():
                    result.append({'type': 'text', 'content': text})
            formula_content = inline_match.group(1)
            if formula_content.strip():
                result.append({'type': 'inline_formula', 'content': '$' + formula_content + '$'})
            pos += inline_match.end()
        else:
            text = content[pos:]
            if text.strip():
                result.append({'type': 'text', 'content': text})
            break
    
    return result

def parse_markdown_tables_from_content(content):
    lines = content.split('\n')
    result_parts = []
    tables = []
    i = 0
    
    while i < len(lines):
        if lines[i].strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) >= 2:
                headers = [h.strip() for h in table_lines[0].split('|') if h.strip()]
                
                if len(table_lines) >= 3 and table_lines[1].strip().startswith('|'):
                    if all(cell.strip() == '' or cell.strip() == '-' * len(cell.strip()) for cell in table_lines[1].split('|')):
                        rows = []
                        for row_line in table_lines[2:]:
                            row = [cell.strip() for cell in row_line.split('|') if cell.strip()]
                            if len(row) == len(headers):
                                rows.append(row)
                        
                        if headers and rows:
                            tables.append({
                                'caption': '',
                                'headers': headers,
                                'rows': rows
                            })
                            result_parts.append({'type': 'table', 'index': len(tables) - 1})
                            continue
            
            result_parts.append({'type': 'text', 'content': '\n'.join(table_lines) + '\n'})
        else:
            result_parts.append({'type': 'text', 'content': lines[i] + '\n'})
            i += 1
    
    return result_parts, tables

def escape_latex(text):
    special_chars = {
        '&': '\\&',
        '%': '\\%',
        '$': '\\$',
        '#': '\\#',
        '_': '\\_',
        '{': '\\{',
        '}': '\\}',
        '~': '\\textasciitilde{}',
        '^': '\\textasciicircum{}',
        '\\': '\\textbackslash{}',
    }
    result = []
    i = 0
    while i < len(text):
        if text[i] in special_chars:
            result.append(special_chars[text[i]])
            i += 1
        else:
            result.append(text[i])
            i += 1
    return ''.join(result)

def _text_to_latex(text):
    text = sanitize_text(text)
    text = text.replace('$$', '\\[')
    text = text.replace('$$', '\\]')
    return text

def generate_paper_latex(title, abstract, sections, image_paths=None):
    title_clean = sanitize_text(title)
    
    abstract_clean = sanitize_text_preserve_newlines(abstract)
    abstract_clean = re.sub(r'\$\$([\s\S]*?)\$\$', r'\\[\1\\]', abstract_clean)
    
    sections_latex = []
    for sec in sections:
        heading = escape_latex(sanitize_text(sec.get("heading", "")))
        content = sanitize_text_preserve_newlines(sec.get("content", ""))
        subsections = sec.get("subsections", [])
        images = sec.get("images", [])
        tables = sec.get("tables", [])
        
        if heading:
            sections_latex.append(f"\\section{{{heading}}}")
        
        if content:
            content_parts, content_tables = parse_markdown_tables_from_content(content)
            
            for part in content_parts:
                if part['type'] == 'table':
                    if part['index'] < len(content_tables):
                        table = content_tables[part['index']]
                        caption = escape_latex(sanitize_text(table.get("caption", "")))
                        headers = [escape_latex(sanitize_text(h)) for h in table.get("headers", [])]
                        rows = [[escape_latex(sanitize_text(cell)) for cell in row] for row in table.get("rows", [])]
                        if headers and rows:
                            col_spec = '|' + '|'.join(['c'] * len(headers)) + '|'
                            table_latex = f"\\begin{{table}}[H]\n\\centering\n\\caption{{{caption}}}\n\\begin{{tabular}}{{{col_spec}}}\n\\hline\n"
                            table_latex += " & ".join(headers) + " \\\\\\hline\n"
                            for row in rows:
                                table_latex += " & ".join(str(cell) for cell in row) + " \\\\\\hline\n"
                            table_latex += "\\end{tabular}\n\\end{table}"
                            sections_latex.append(table_latex)
                elif part['type'] == 'text':
                    text_parts = split_content_with_formulas(part['content'])
                    latex_text = ""
                    for text_part in text_parts:
                        if text_part['type'] == 'text':
                            latex_text += escape_latex(text_part['content'])
                        elif text_part['type'] == 'inline_formula':
                            latex_text += text_part['content']
                        elif text_part['type'] == 'block_formula':
                            formula_content = text_part['content'][2:-2]
                            latex_text += f"\\begin{{equation}}\n{formula_content}\n\\end{{equation}}\n"
                    if latex_text.strip():
                        sections_latex.append(latex_text)
        
        for subsection in subsections:
            sub_heading = escape_latex(sanitize_text(subsection.get("heading", "")))
            sub_content = sanitize_text_preserve_newlines(subsection.get("content", ""))
            
            if sub_heading:
                sections_latex.append(f"\\subsection{{{sub_heading}}}")
            
            if sub_content:
                sub_content_parts, sub_content_tables = parse_markdown_tables_from_content(sub_content)
                
                for part in sub_content_parts:
                    if part['type'] == 'table':
                        if part['index'] < len(sub_content_tables):
                            table = sub_content_tables[part['index']]
                            caption = escape_latex(sanitize_text(table.get("caption", "")))
                            headers = [escape_latex(sanitize_text(h)) for h in table.get("headers", [])]
                            rows = [[escape_latex(sanitize_text(cell)) for cell in row] for row in table.get("rows", [])]
                            if headers and rows:
                                col_spec = '|' + '|'.join(['c'] * len(headers)) + '|'
                                table_latex = f"\\begin{{table}}[H]\n\\centering\n\\caption{{{caption}}}\n\\begin{{tabular}}{{{col_spec}}}\n\\hline\n"
                                table_latex += " & ".join(headers) + " \\\\\\hline\n"
                                for row in rows:
                                    table_latex += " & ".join(str(cell) for cell in row) + " \\\\\\hline\n"
                                table_latex += "\\end{tabular}\n\\end{table}"
                                sections_latex.append(table_latex)
                    elif part['type'] == 'text':
                        text_parts = split_content_with_formulas(part['content'])
                        latex_text = ""
                        for text_part in text_parts:
                            if text_part['type'] == 'text':
                                latex_text += escape_latex(text_part['content'])
                            elif text_part['type'] == 'inline_formula':
                                latex_text += text_part['content']
                            elif text_part['type'] == 'block_formula':
                                formula_content = text_part['content'][2:-2]
                                latex_text += f"\\begin{{equation}}\n{formula_content}\n\\end{{equation}}\n"
                        if latex_text.strip():
                            sections_latex.append(latex_text)
        
        for table in tables:
            caption = escape_latex(sanitize_text(table.get("caption", "")))
            headers = [escape_latex(sanitize_text(h)) for h in table.get("headers", [])]
            rows = [[escape_latex(sanitize_text(cell)) for cell in row] for row in table.get("rows", [])]
            if headers and rows:
                col_spec = '|' + '|'.join(['c'] * len(headers)) + '|'
                table_latex = f"\\begin{{table}}[H]\n\\centering\n\\caption{{{caption}}}\n\\begin{{tabular}}{{{col_spec}}}\n\\hline\n"
                table_latex += " & ".join(headers) + " \\\\\\hline\n"
                for row in rows:
                    table_latex += " & ".join(str(cell) for cell in row) + " \\\\\\hline\n"
                table_latex += "\\end{tabular}\n\\end{table}"
                sections_latex.append(table_latex)
        
        for img in images:
            img_path = img.get("path", "")
            img_caption = escape_latex(sanitize_text(img.get("caption", "")))
            if img_path:
                filename = os.path.basename(img_path)
                sections_latex.append(f"\\begin{{figure}}[H]\n\\centering\n\\includegraphics[width=0.8\\textwidth]{{{filename}}}\n\\caption{{{img_caption}}}\n\\end{{figure}}")
    
    title_escaped = escape_latex(title_clean)
    
    template = f"""\\documentclass[12pt,a4paper]{{ctexart}}
\\usepackage{{amsmath,amssymb,amsfonts}}
\\usepackage{{graphicx}}
\\usepackage{{booktabs}}
\\usepackage{{geometry}}
\\geometry{{left=2.5cm,right=2.5cm,top=2.5cm,bottom=2.5cm}}
\\usepackage{{hyperref}}
\\usepackage{{float}}
\\usepackage{{caption}}

\\title{{\\textbf{{{title_escaped if title_escaped else '数学建模竞赛论文'}}}}}
\\author{{}}
\\date{{}}

\\begin{{document}}
\\maketitle

\\begin{{abstract}}
{abstract_clean}
\\end{{abstract}}

\\tableofcontents
\\newpage

{chr(10).join(sections_latex)}

\\end{{document}}
"""
    return template

def add_paragraph_with_format(doc, text, font_name='宋体', font_size=12, bold=False, first_line_indent=True, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    
    if first_line_indent:
        para_format = para.paragraph_format
        para_format.first_line_indent = Cm(0.74)
    
    para.alignment = alignment
    return para

def generate_paper_docx(title, abstract, sections, output_path, image_paths=None):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    para_format = style.paragraph_format
    para_format.line_spacing = 1.5
    para_format.first_line_indent = Cm(0.74)
    
    title_clean = sanitize_text(title) if title else '数学建模竞赛论文'
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_clean)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(22)
    title_run.bold = True
    
    doc.add_paragraph()
    
    abstract_heading = doc.add_paragraph()
    abstract_run = abstract_heading.add_run('摘要')
    abstract_run.font.name = '黑体'
    abstract_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    abstract_run.font.size = Pt(16)
    abstract_run.bold = True
    
    abstract_clean = sanitize_text_preserve_newlines(abstract) if abstract else '本论文针对数学建模问题进行了深入分析和研究...'
    abstract_parts = split_content_with_formulas(abstract_clean)
    if abstract_parts:
        abstract_para = doc.add_paragraph()
        abstract_para.paragraph_format.first_line_indent = Cm(0.74)
        for part in abstract_parts:
            if part['type'] == 'text':
                run = abstract_para.add_run(part['content'])
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
            elif part['type'] == 'inline_formula':
                add_latex_formula(abstract_para, part['content'])
            elif part['type'] == 'block_formula':
                doc.add_paragraph()
                formula_para = doc.add_paragraph()
                formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_latex_formula(formula_para, part['content'])
                doc.add_paragraph()
                abstract_para = doc.add_paragraph()
                abstract_para.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph()
    
    output_dir = os.path.dirname(output_path)
    
    for sec in sections:
        heading = sanitize_text(sec.get("heading", ""))
        content = sanitize_text_preserve_newlines(sec.get("content", ""))
        subsections = sec.get("subsections", [])
        images = sec.get("images", [])
        tables = sec.get("tables", [])
        
        if heading:
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(heading)
            heading_run.font.name = '黑体'
            heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            heading_run.font.size = Pt(16)
            heading_run.bold = True
        
        if content:
            content_parts, content_tables = parse_markdown_tables_from_content(content)
            
            for part in content_parts:
                if part['type'] == 'table':
                    if part['index'] < len(content_tables):
                        add_table_to_docx(doc, content_tables[part['index']])
                elif part['type'] == 'text':
                    text_parts = split_content_with_formulas(part['content'])
                    if text_parts:
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Cm(0.74)
                        for text_part in text_parts:
                            if text_part['type'] == 'text':
                                run = para.add_run(text_part['content'])
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                run.font.size = Pt(12)
                            elif text_part['type'] == 'inline_formula':
                                add_latex_formula(para, text_part['content'])
                            elif text_part['type'] == 'block_formula':
                                doc.add_paragraph()
                                formula_para = doc.add_paragraph()
                                formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                add_latex_formula(formula_para, text_part['content'])
                                doc.add_paragraph()
                                para = doc.add_paragraph()
                                para.paragraph_format.first_line_indent = Cm(0.74)
        
        if subsections:
            for subsection in subsections:
                sub_heading = sanitize_text(subsection.get("heading", ""))
                sub_content = sanitize_text_preserve_newlines(subsection.get("content", ""))
                sub_images = subsection.get("images", [])
                sub_tables = subsection.get("tables", [])
                
                if sub_heading:
                    sub_heading_para = doc.add_paragraph()
                    sub_heading_run = sub_heading_para.add_run(sub_heading)
                    sub_heading_run.font.name = '黑体'
                    sub_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    sub_heading_run.font.size = Pt(14)
                    sub_heading_run.bold = True
                
                if sub_content:
                    sub_content_parts, sub_content_tables = parse_markdown_tables_from_content(sub_content)
                    
                    for part in sub_content_parts:
                        if part['type'] == 'table':
                            if part['index'] < len(sub_content_tables):
                                add_table_to_docx(doc, sub_content_tables[part['index']])
                        elif part['type'] == 'text':
                            text_parts = split_content_with_formulas(part['content'])
                            if text_parts:
                                para = doc.add_paragraph()
                                para.paragraph_format.first_line_indent = Cm(0.74)
                                for text_part in text_parts:
                                    if text_part['type'] == 'text':
                                        run = para.add_run(text_part['content'])
                                        run.font.name = '宋体'
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                        run.font.size = Pt(12)
                                    elif text_part['type'] == 'inline_formula':
                                        add_latex_formula(para, text_part['content'])
                                    elif text_part['type'] == 'block_formula':
                                        doc.add_paragraph()
                                        formula_para = doc.add_paragraph()
                                        formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        add_latex_formula(formula_para, text_part['content'])
                                        doc.add_paragraph()
                                        para = doc.add_paragraph()
                                        para.paragraph_format.first_line_indent = Cm(0.74)
                
                if sub_tables:
                    for table in sub_tables:
                        add_table_to_docx(doc, table)
                
                if sub_images:
                    for img in sub_images:
                        img_path = img.get("path", "")
                        img_caption = sanitize_text(img.get("caption", ""))
                        
                        if not img_path:
                            continue
                        
                        if os.path.exists(img_path):
                            pass
                        elif output_dir and os.path.exists(os.path.join(output_dir, os.path.basename(img_path))):
                            img_path = os.path.join(output_dir, os.path.basename(img_path))
                        else:
                            continue
                        
                        try:
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_para.add_run()
                            run.add_picture(img_path, width=Inches(5))
                            if img_caption:
                                caption_para = doc.add_paragraph(img_caption)
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption_run = caption_para.add_run()
                                caption_run.font.name = '宋体'
                                caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                caption_run.font.size = Pt(10)
                            doc.add_paragraph()
                        except Exception as e:
                            doc.add_paragraph(f"图片插入失败: {os.path.basename(img_path)}")
        
        if tables:
            for table in tables:
                add_table_to_docx(doc, table)
        
        if images:
            for img in images:
                img_path = img.get("path", "")
                img_caption = sanitize_text(img.get("caption", ""))
                
                if not img_path:
                    continue
                
                if os.path.exists(img_path):
                    pass
                elif output_dir and os.path.exists(os.path.join(output_dir, os.path.basename(img_path))):
                    img_path = os.path.join(output_dir, os.path.basename(img_path))
                else:
                    continue
                
                try:
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_para.add_run()
                    run.add_picture(img_path, width=Inches(5))
                    if img_caption:
                        caption_para = doc.add_paragraph(img_caption)
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_para.add_run()
                        caption_run.font.name = '宋体'
                        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        caption_run.font.size = Pt(10)
                    doc.add_paragraph()
                except Exception as e:
                    doc.add_paragraph(f"图片插入失败: {os.path.basename(img_path)}")
    
    doc.save(output_path)
    return output_path

def sanitize_text_preserve_newlines(text):
    if text is None:
        return ""
    text = str(text)
    text = text.replace('\x00', '')
    text = text.replace('\r', '')
    text = text.replace('\t', ' ')
    text = re.sub(r'[\x00-\x0B\x0E-\x1F\x7F]', '', text)
    return text.strip()

def generate_paper_markdown(title, abstract, sections, image_paths=None):
    title_clean = sanitize_text(title) if title else '数学建模竞赛论文'
    abstract_clean = sanitize_text_preserve_newlines(abstract) if abstract else '本论文针对数学建模问题进行了深入分析和研究...'
    
    md = f"# {title_clean}\n\n"
    md += f"## 摘要\n\n{abstract_clean}\n\n"
    md += "---\n\n"
    
    for sec_idx, sec in enumerate(sections):
        heading = sanitize_text(sec.get("heading", ""))
        content = sanitize_text_preserve_newlines(sec.get("content", ""))
        subsections = sec.get("subsections", [])
        images = sec.get("images", [])
        tables = sec.get("tables", [])
        
        if heading:
            md += f"## {heading}\n\n"
        
        if content:
            md += f"{content}\n\n"
        
        if subsections:
            for sub_idx, subsection in enumerate(subsections):
                sub_heading = sanitize_text(subsection.get("heading", ""))
                sub_content = sanitize_text_preserve_newlines(subsection.get("content", ""))
                sub_images = subsection.get("images", [])
                sub_tables = subsection.get("tables", [])
                
                if sub_heading:
                    md += f"### {sub_heading}\n\n"
                
                if sub_content:
                    md += f"{sub_content}\n\n"
                
                if sub_tables:
                    for table in sub_tables:
                        caption = sanitize_text(table.get("caption", ""))
                        headers = [sanitize_text(h) for h in table.get("headers", [])]
                        rows = [[sanitize_text(cell) for cell in row] for row in table.get("rows", [])]
                        if caption:
                            md += f"**{caption}**\n\n"
                        if headers and rows:
                            md += "| " + " | ".join(headers) + " |\n"
                            md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
                            for row in rows:
                                md += "| " + " | ".join(str(cell) for cell in row) + " |\n"
                            md += "\n"
                
                if sub_images:
                    for img in sub_images:
                        img_path = img.get("path", "")
                        img_caption = sanitize_text(img.get("caption", ""))
                        if img_path:
                            filename = os.path.basename(img_path)
                            md += f"![{img_caption}]({filename})\n\n"
        
        if tables:
            for table in tables:
                caption = sanitize_text(table.get("caption", ""))
                headers = [sanitize_text(h) for h in table.get("headers", [])]
                rows = [[sanitize_text(cell) for cell in row] for row in table.get("rows", [])]
                if caption:
                    md += f"**{caption}**\n\n"
                if headers and rows:
                    md += "| " + " | ".join(headers) + " |\n"
                    md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
                    for row in rows:
                        md += "| " + " | ".join(str(cell) for cell in row) + " |\n"
                    md += "\n"
        
        if images:
            for img in images:
                img_path = img.get("path", "")
                img_caption = sanitize_text(img.get("caption", ""))
                if img_path:
                    filename = os.path.basename(img_path)
                    md += f"![{img_caption}]({filename})\n\n"
        
        if sec_idx < len(sections) - 1:
            md += "---\n\n"
    
    return md

def convert_markdown_to_docx(md_file_path, output_docx_path=None, style_path=None):
    if not os.path.exists(md_file_path):
        raise FileNotFoundError(f"Markdown文件不存在: {md_file_path}")
    
    if not os.path.exists(MARKDOCX_PATH):
        raise FileNotFoundError(f"markdocx工具路径不存在: {MARKDOCX_PATH}")
    
    if output_docx_path is None:
        output_docx_path = md_file_path.replace('.md', '.docx')
    
    md_file_path_abs = os.path.abspath(md_file_path)
    output_docx_path_abs = os.path.abspath(output_docx_path)
    
    if style_path is None:
        style_path = os.path.join(os.path.dirname(__file__), 'paper_style.yaml')
    
    style_path_abs = os.path.abspath(style_path)
    
    original_cwd = os.getcwd()
    
    try:
        os.chdir(MARKDOCX_PATH)
        
        cmd = [
            sys.executable, 'src/markdocx.py',
            md_file_path_abs,
            '-o', output_docx_path_abs
        ]
        
        if os.path.exists(style_path_abs):
            cmd.extend(['-s', style_path_abs])
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode == 0 or 'SUCCESS' in result.stdout:
            print(f"[SUCCESS] Markdown转DOCX成功: {output_docx_path_abs}")
            return output_docx_path_abs
        else:
            print(f"[ERROR] Markdown转DOCX失败: {result.stderr[:200]}")
            print(f"[WARNING] 将使用内置docx库生成DOCX")
            return None
    except subprocess.TimeoutExpired:
        print(f"[ERROR] Markdown转DOCX超时")
        print(f"[WARNING] 将使用内置docx库生成DOCX")
        return None
    except Exception as e:
        print(f"[ERROR] Markdown转DOCX异常: {e}")
        print(f"[WARNING] 将使用内置docx库生成DOCX")
        return None
    finally:
        os.chdir(original_cwd)

def generate_paper(title, abstract, sections, format='markdown', output_dir=None):
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'output')
    os.makedirs(output_dir, exist_ok=True)
    
    title_clean = sanitize_text(title) if title else '数学建模竞赛论文'
    title_clean = re.sub(r'[\\/*?:"<>|]', '', title_clean)
    if len(title_clean) > 50:
        title_clean = title_clean[:50]
    
    md_content = generate_paper_markdown(title_clean, abstract, sections)
    md_output_path = os.path.join(output_dir, f'{title_clean}_论文.md')
    
    with open(md_output_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    if format == 'markdown':
        return md_output_path
    elif format == 'docx':
        docx_path = os.path.join(output_dir, f'{title_clean}_论文.docx')
        markdocx_result = convert_markdown_to_docx(md_output_path, docx_path)
        if markdocx_result:
            return markdocx_result
        else:
            return generate_paper_docx(title_clean, abstract, sections, docx_path)
    elif format == 'latex':
        output_path = os.path.join(output_dir, f'{title_clean}_论文.tex')
        latex_content = generate_paper_latex(title_clean, abstract, sections)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        return output_path
    else:
        raise ValueError(f"不支持的格式: {format}")