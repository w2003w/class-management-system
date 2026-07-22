import io
import os
import json
from pathlib import Path

import pandas as pd
import numpy as np

def analyze_data_file(file_path):
    try:
        ext = Path(file_path).suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(file_path, encoding="utf-8")
        elif ext in (".xls", ".xlsx"):
            df = pd.read_excel(file_path)
        else:
            return f"不支持的文件格式: {ext}，仅支持 .csv / .xls / .xlsx"

        buf = io.StringIO()
        buf.write(f"📊 文件: {Path(file_path).name}\n")
        buf.write(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n")
        buf.write(f"总行数: {df.shape[0]}  |  总列数: {df.shape[1]}\n")
        buf.write(f"列名: {list(df.columns)}\n")
        buf.write(f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n")
        buf.write(f"完整数据内容:\n")
        buf.write(df.to_csv(index=False, encoding='utf-8'))
        return buf.getvalue()
    except Exception as e:
        return f"❌ 文件读取失败: {e}"

def clean_data(df):
    cleaned = df.copy()
    for col in cleaned.columns:
        if cleaned[col].dtype in ['float64', 'int64']:
            median_val = cleaned[col].median()
            cleaned[col].fillna(median_val, inplace=True)
            z_scores = np.abs((cleaned[col] - cleaned[col].mean()) / cleaned[col].std())
            cleaned = cleaned[z_scores < 3]
    return cleaned

def parse_pdf(file_path):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        lines = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
        return '\n'.join(lines)
    except ImportError:
        return "[PyPDF2 未安装]"

def parse_docx(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(' | '.join(cells))
        return '\n'.join(lines)
    except ImportError:
        return "[python-docx 未安装]"

def clean_text(text):
    if text is None:
        return ""
    
    cleaned = []
    for char in text:
        try:
            char.encode('utf-8')
            cleaned.append(char)
        except UnicodeEncodeError:
            pass
    
    return ''.join(cleaned)

def parse_file(file_path):
    ext = Path(file_path).suffix.lower()[1:]
    
    if ext == 'pdf':
        return clean_text(parse_pdf(file_path))
    elif ext == 'docx':
        return clean_text(parse_docx(file_path))
    elif ext in ('xlsx', 'xls'):
        return clean_text(analyze_data_file(file_path))
    elif ext == 'csv':
        return clean_text(analyze_data_file(file_path))
    elif ext in ('txt', 'md'):
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    return clean_text(f.read())
            except UnicodeDecodeError:
                continue
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return clean_text(f.read())
    else:
        return f"[不支持的文件类型: {ext}]"

def generate_eda_charts(df, output_dir):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import seaborn as sns
    
    os.makedirs(output_dir, exist_ok=True)
    plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    
    numeric_cols = df.select_dtypes(include=['float64', 'int64']).columns
    
    if len(numeric_cols) > 0:
        plt.figure(figsize=(10, 6))
        sns.boxplot(data=df[numeric_cols])
        plt.title('数值列箱线图')
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, 'boxplot.png'), dpi=100)
        plt.close()
    
    if len(numeric_cols) >= 2:
        corr = df[numeric_cols].corr()
        plt.figure(figsize=(8, 6))
        sns.heatmap(corr, annot=True, cmap='coolwarm', fmt='.2f')
        plt.title('相关性热力图')
        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, 'correlation.png'), dpi=100)
        plt.close()
    
    for col in numeric_cols[:5]:
        plt.figure(figsize=(8, 4))
        sns.histplot(df[col], kde=True)
        plt.title(f'{col} 分布')
        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, f'{col}_hist.png'), dpi=100)
        plt.close()
    
    return output_dir